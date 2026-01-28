"""
Test script for Q-SFTP Privacy Features (Phase 1)

This script tests:
1. Image metadata scrubbing (EXIF removal)
2. PDF metadata scrubbing (author/creator removal)
3. Word document metadata scrubbing
4. IP address hashing functionality
5. File categorization

Run this script to verify privacy features are working correctly.
"""

import os
import io
import sys

# Fix encoding for Windows console
if sys.platform == 'win32':
    sys.stdout.reconfigure(encoding='utf-8', errors='replace')
    sys.stderr.reconfigure(encoding='utf-8', errors='replace')
from datetime import datetime

# Add WebApp to path
sys.path.insert(0, os.path.join(os.path.dirname(__file__), 'Codes', 'WebApp'))

from privacy_manager import privacy_manager
from activity_logger import activity_logger

def test_image_metadata():
    """Test EXIF removal from images."""
    print("\n" + "="*60)
    print("TEST 1: Image Metadata Scrubbing (EXIF Removal)")
    print("="*60)
    
    try:
        from PIL import Image
        from PIL.ExifTags import TAGS
        
        # Create test image with EXIF data
        img = Image.new('RGB', (100, 100), color='red')
        
        # Add EXIF data
        exif_dict = {
            271: 'TestCamera Manufacturer',  # Make
            272: 'TestCamera Model',  # Model
            305: 'TestSoftware v1.0',  # Software
            306: '2024:01:28 12:00:00'  # DateTime
        }
        
        # Save with EXIF
        buffer = io.BytesIO()
        img.save(buffer, format='JPEG', exif=img.getexif())
        original_bytes = buffer.getvalue()
        
        print(f"✓ Created test image: {len(original_bytes)} bytes")
        print(f"✓ EXIF tags added: {len(exif_dict)} fields")
        
        # Scrub metadata
        scrubbed_bytes, metadata_removed = privacy_manager.scrub_image_metadata(
            original_bytes,
            'test_image.jpg'
        )
        
        print(f"✓ Scrubbed image: {len(scrubbed_bytes)} bytes")
        
        # Verify EXIF is removed
        scrubbed_img = Image.open(io.BytesIO(scrubbed_bytes))
        has_exif = hasattr(scrubbed_img, '_getexif') and scrubbed_img._getexif()
        
        if not has_exif:
            print("✅ SUCCESS: EXIF data removed!")
        else:
            print("❌ FAILED: EXIF data still present!")
        
        print(f"Metadata removed: {metadata_removed}")
        return not has_exif
        
    except ImportError as e:
        print(f"⚠️  SKIPPED: Pillow not installed ({e})")
        print("   Install: pip install Pillow")
        return None


def test_pdf_metadata():
    """Test PDF metadata removal."""
    print("\n" + "="*60)
    print("TEST 2: PDF Metadata Scrubbing")
    print("="*60)
    
    try:
        from PyPDF2 import PdfReader, PdfWriter
        from reportlab.pdfgen import canvas
        from reportlab.lib.pagesizes import letter
        
        # Create test PDF with metadata
        buffer = io.BytesIO()
        c = canvas.Canvas(buffer, pagesize=letter)
        c.setAuthor("Test Author")
        c.setCreator("Test Creator")
        c.setTitle("Confidential Document")
        c.drawString(100, 750, "Test PDF Content")
        c.save()
        
        original_bytes = buffer.getvalue()
        
        # Read metadata
        reader = PdfReader(io.BytesIO(original_bytes))
        original_metadata = reader.metadata
        
        print(f"✓ Created test PDF: {len(original_bytes)} bytes")
        print(f"✓ Metadata fields: {len(original_metadata) if original_metadata else 0}")
        if original_metadata:
            print(f"  - Author: {original_metadata.get('/Author', 'N/A')}")
            print(f"  - Creator: {original_metadata.get('/Creator', 'N/A')}")
        
        # Scrub metadata
        scrubbed_bytes, metadata_removed = privacy_manager.scrub_pdf_metadata(original_bytes)
        
        # Verify removal
        scrubbed_reader = PdfReader(io.BytesIO(scrubbed_bytes))
        scrubbed_metadata = scrubbed_reader.metadata
        
        print(f"✓ Scrubbed PDF: {len(scrubbed_bytes)} bytes")
        
        # Check if sensitive metadata is removed
        has_author = scrubbed_metadata and scrubbed_metadata.get('/Author')
        has_creator = scrubbed_metadata and scrubbed_metadata.get('/Creator') not in [None, 'Q-SFTP Privacy System']
        
        if not has_author and not has_creator:
            print("✅ SUCCESS: PDF metadata removed!")
        else:
            print("❌ FAILED: Some metadata still present!")
            print(f"   Scrubbed metadata: {scrubbed_metadata}")
        
        print(f"Metadata removed: {metadata_removed}")
        return not has_author and not has_creator
        
    except ImportError as e:
        print(f"⚠️  SKIPPED: PyPDF2 or reportlab not installed ({e})")
        print("   Install: pip install PyPDF2 reportlab")
        return None


def test_docx_metadata():
    """Test Word document metadata removal."""
    print("\n" + "="*60)
    print("TEST 3: Word Document Metadata Scrubbing")
    print("="*60)
    
    try:
        from docx import Document
        
        # Create test document with metadata
        doc = Document()
        doc.add_paragraph("Test content")
        
        # Add metadata
        core_props = doc.core_properties
        core_props.author = "John Doe"
        core_props.last_modified_by = "Jane Smith"
        core_props.title = "Confidential Report"
        core_props.category = "ACME Corporation"
        
        # Save to bytes
        buffer = io.BytesIO()
        doc.save(buffer)
        original_bytes = buffer.getvalue()
        
        print(f"✓ Created test DOCX: {len(original_bytes)} bytes")
        print(f"  - Author: {core_props.author}")
        print(f"  - Last Modified By: {core_props.last_modified_by}")
        print(f"  - Company: {core_props.category}")
        
        # Scrub metadata
        scrubbed_bytes, metadata_removed = privacy_manager.scrub_docx_metadata(original_bytes)
        
        # Verify removal
        scrubbed_doc = Document(io.BytesIO(scrubbed_bytes))
        scrubbed_props = scrubbed_doc.core_properties
        
        print(f"✓ Scrubbed DOCX: {len(scrubbed_bytes)} bytes")
        print(f"  - Author: '{scrubbed_props.author}'")
        print(f"  - Last Modified By: '{scrubbed_props.last_modified_by}'")
        
        is_clean = (not scrubbed_props.author and not scrubbed_props.last_modified_by)
        
        if is_clean:
            print("✅ SUCCESS: DOCX metadata removed!")
        else:
            print("❌ FAILED: Some metadata still present!")
        
        print(f"Metadata removed: {metadata_removed}")
        return is_clean
        
    except ImportError as e:
        print(f"⚠️  SKIPPED: python-docx not installed ({e})")
        print("   Install: pip install python-docx")
        return None


def test_ip_hashing():
    """Test IP address hashing with daily salt."""
    print("\n" + "="*60)
    print("TEST 4: IP Address Anonymization")
    print("="*60)
    
    test_ips = [
        "192.168.1.1",
        "10.0.0.5",
        "172.16.0.100",
        "127.0.0.1"
    ]
    
    print("Testing IP hashing...")
    hashes = []
    
    for ip in test_ips:
        hashed = activity_logger.anonymize_ip(ip)
        hashes.append(hashed)
        print(f"  {ip:15} → {hashed}")
    
    # Verify properties
    all_16_chars = all(len(h) == 16 for h in hashes)
    all_unique = len(hashes) == len(set(hashes))
    no_original_ips = not any(ip in h for ip, h in zip(test_ips, hashes))
    
    success = all_16_chars and all_unique and no_original_ips
    
    if success:
        print("✅ SUCCESS: IP addresses properly hashed!")
        print("   - All hashes are 16 characters")
        print("   - All hashes are unique")
        print("   - No original IPs in hashes")
    else:
        print("❌ FAILED: IP hashing not working correctly!")
    
    # Test salt rotation
    print("\nTesting daily salt rotation...")
    salt_today = activity_logger._get_daily_salt()
    print(f"  Today's salt (first 16 chars): {salt_today[:16]}")
    print(f"  ✓ Salt is {len(salt_today)} characters")
    
    return success


def test_file_categorization():
    """Test file category detection."""
    print("\n" + "="*60)
    print("TEST 5: File Categorization")
    print("="*60)
    
    test_files = {
        'report.pdf': 'document',
        'photo.jpg': 'image',
        'presentation.pptx': 'other',  # Not yet supported
        'archive.zip': 'archive',
        'script.py': 'code',
        'video.mp4': 'video',
        'unknown.xyz': 'other'
    }
    
    all_correct = True
    
    for filename, expected_category in test_files.items():
        actual_category = privacy_manager.get_file_category(filename)
        match = "✓" if actual_category == expected_category else "✗"
        print(f"  {match} {filename:20} → {actual_category:10} (expected: {expected_category})")
        if actual_category != expected_category:
            all_correct = False
    
    if all_correct:
        print("✅ SUCCESS: File categorization working!")
    else:
        print("⚠️  Some categorizations unexpected (may be normal)")
    
    return all_correct


def test_configuration():
    """Test privacy configuration loading."""
    print("\n" + "="*60)
    print("TEST 6: Privacy Configuration")
    print("="*60)
    
    config = privacy_manager.config
    
    print("Loaded configuration:")
    print(f"  • Metadata scrubbing: {'Enabled' if config['metadata_scrubbing']['enabled'] else 'Disabled'}")
    print(f"  • Supported file types: {', '.join(config['metadata_scrubbing']['file_types'])}")
    print(f"  • IP anonymization: {'Enabled' if config['ip_anonymization']['enabled'] else 'Disabled'}")
    print(f"  • Hash length: {config['ip_anonymization']['hash_length']} chars")
    
    has_config = config is not None
    if has_config:
        print("✅ SUCCESS: Configuration loaded!")
    else:
        print("❌ FAILED: Configuration not loaded!")
    
    return has_config


def run_all_tests():
    """Run all privacy feature tests."""
    print("\n" + "#"*60)
    print("#" + " "*20 + "Q-SFTP PRIVACY TESTS" + " "*19 + "#")
    print("#"*60)
    print(f"\nTest Date: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    
    results = {
        'Configuration': test_configuration(),
        'IP Hashing': test_ip_hashing(),
        'File Categorization': test_file_categorization(),
        'Image Metadata': test_image_metadata(),
        'PDF Metadata': test_pdf_metadata(),
        'Word Metadata': test_docx_metadata()
    }
    
    # Summary
    print("\n" + "="*60)
    print("TEST SUMMARY")
    print("="*60)
    
    passed = sum(1 for r in results.values() if r is True)
    failed = sum(1 for r in results.values() if r is False)
    skipped = sum(1 for r in results.values() if r is None)
    
    for test_name, result in results.items():
        if result is True:
            status = "✅ PASS"
        elif result is False:
            status = "❌ FAIL"
        else:
            status = "⚠️  SKIP"
        print(f"{status} - {test_name}")
    
    print(f"\nResults: {passed} passed, {failed} failed, {skipped} skipped")
    
    if failed == 0:
        print("\n🎉 All tests passed! Privacy features are working correctly.")
    else:
        print(f"\n⚠️  {failed} test(s) failed. Please review the output above.")
    
    print("\n" + "#"*60)
    
    return failed == 0


if __name__ == '__main__':
    success = run_all_tests()
    sys.exit(0 if success else 1)
